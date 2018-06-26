import sys

reload(sys)
sys.setdefaultencoding('utf-8')
import textactor
import subprocess
import docx
import os
import urllib

from docx import Document
from PyPDF2.pdf import PdfFileReader
# from pyPdf import PdfFileReader
from cStringIO import StringIO

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage

import urllib2


def convert_pdf_to_text(link, path):
    # urllib.urlretrieve(link, path)
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    # device = TextConverter(rsrcmgr, retstr, codec='utf-8', laparams=LAParams())
    device = TextConverter(rsrcmgr, retstr, laparams=LAParams())
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    with open(path, 'rb') as fp:
        for page in PDFPage.get_pages(fp, set()):
            interpreter.process_page(page)
        text = retstr.getvalue()

    device.close()
    retstr.close()
    with open(path[:(path.index('.'))] + '.text', 'wb') as f:
        f.write(text)


def pdf_read(file_name):
    # input1 = PdfFileReader(file(file_name, 'rb'))
    # # print "title = %s" % (input1.getDocumentInfo().title)
    # # print input1.pages
    # for page in input1.getPage(1):
    #     print page.extractText()
    pdf = PdfFileReader(file(file_name, "rb"))
    content = ""
    pagecount = pdf.getNumPages()
    print('pagecount:', pagecount)
    # pageLabels = {}
    # page = pdf.getPage(0)
    # pageLabels[page.indirectRef.idnum] =  1
    # print pdf.getOutlines()[0]['/Page']['/Contents']
    for i in range(0, pdf.getNumPages()):
        extractedText = pdf.getPage(i).extractText()
        content += extractedText + "\n"
        # return content.encode("ascii", "ignore")
    print content


def word(path):
    for filename in os.listdir(os.getcwd()):
        if filename.endswith('.doc'):
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filename])
            doc = docx.Document(filename[:-4] + ".docx")
            for para in doc.paragraphs:
                print (para.text)
                # subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filename])
                # doc = docx.Document(filename[:-4] + ".docx")
                # for para in doc.paragraphs:
                #     print (para.text)


def get_response(url, file_name):
    # response = urllib2.urlopen(url)
    urllib.urlretrieve(url, file_name)

    # if file_name.endswith('pdf'):
    #     convert_pdf_to_text(file_name)


    # if not os.path.exists('test.docx'):
    #     os.mknod('test.docx')
    # with open('test.docx','wb') as f:
    #     f.write(response.read())

    return file_name


def main():
    url = 'http://download.zikao365.com/shiti/18766.doc'
    pdf_url = "http://download.zikao365.com/shiti/35247.pdf"

    filename = "a.pdf"

    # get_response(pdf_url, filename)
    # path = get_word(url)
    pdf_read(filename)

    # word(path)
    # from subprocess import call
    # call('qpdf --password=%s --decrypt %s %s' % ('', filename, filename), shell=True)
    # convert_pdf_to_text(pdf_url,filename)


if __name__ == "__main__":
    main()
